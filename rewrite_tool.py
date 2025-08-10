from crewai.tools import BaseTool
from typing import Dict, Any, List
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import json
from datetime import datetime

class DocumentRewriterTool(BaseTool):
    name: str = "ADGM Document Rewriter Tool"
    description: str = "Rewrites ADGM documents with compliance comments and generates edit reports"
    
    def _run(self, rewrite_instructions: List[Dict] = None) -> Dict[str, Any]:
        """
        Rewrites documents with compliance comments and tracks changes
        Args:
            rewrite_instructions: List of {
                "filename": "doc.docx",
                "violations": [{"text": "old text", "replacement": "new text", "comment": "compliance issue", "severity": "HIGH"}]
            }
        """
        
        if not rewrite_instructions:
            return {"error": "No rewrite instructions provided", "status": "error"}
        
        documents_dir = os.path.join(os.getcwd(), "documents")
        outputs_dir = os.path.join(os.getcwd(), "corrected_documents")
        
        os.makedirs(outputs_dir, exist_ok=True)
        
        edit_report = {
            "timestamp": datetime.now().isoformat(),
            "documents_processed": [],
            "total_edits": 0,
            "critical_fixes": 0,
            "high_fixes": 0,
            "medium_fixes": 0,
            "low_fixes": 0
        }
        
        for instruction in rewrite_instructions:
            filename = instruction.get("filename")
            violations = instruction.get("violations", [])
            
            if not filename or not violations:
                continue
                
            result = self._process_document(filename, violations, documents_dir, outputs_dir)
            edit_report["documents_processed"].append(result)
            
            edit_report["total_edits"] += result["edits_made"]
            edit_report["critical_fixes"] += result["severity_counts"]["CRITICAL"]
            edit_report["high_fixes"] += result["severity_counts"]["HIGH"] 
            edit_report["medium_fixes"] += result["severity_counts"]["MEDIUM"]
            edit_report["low_fixes"] += result["severity_counts"]["LOW"]
        
        report_path = os.path.join(outputs_dir, "compliance_edit_report.json")
        with open(report_path, 'w') as f:
            json.dump(edit_report, f, indent=2)
        
        return {
            "status": "success",
            "documents_processed": len(rewrite_instructions),
            "total_edits": edit_report["total_edits"],
            "report_path": report_path,
            "corrected_documents_dir": outputs_dir,
            "edit_summary": edit_report
        }
    
    def _process_document(self, filename: str, violations: List[Dict], input_dir: str, output_dir: str) -> Dict:
        """Process individual document with violations"""
        
        input_path = os.path.join(input_dir, filename)
        output_filename = f"CORRECTED_{filename}"
        output_path = os.path.join(output_dir, output_filename)
        
        doc_result = {
            "original_filename": filename,
            "corrected_filename": output_filename,
            "edits_made": 0,
            "severity_counts": {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0},
            "violations_fixed": [],
            "status": "success"
        }
        
        try:
            doc = Document(input_path)
            
            for violation in violations:
                old_text = violation.get("text", "")
                new_text = violation.get("replacement", "")
                comment = violation.get("comment", "")
                severity = violation.get("severity", "LOW")
                
                if self._replace_text_with_comment(doc, old_text, new_text, comment, severity):
                    doc_result["edits_made"] += 1
                    doc_result["severity_counts"][severity] += 1
                    doc_result["violations_fixed"].append({
                        "original_text": old_text[:50] + "..." if len(old_text) > 50 else old_text,
                        "corrected_text": new_text[:50] + "..." if len(new_text) > 50 else new_text,
                        "severity": severity,
                        "comment": comment
                    })
            
            self._add_compliance_header(doc, doc_result["edits_made"], filename)
            
            doc.save(output_path)
            
        except Exception as e:
            doc_result["status"] = "error"
            doc_result["error"] = str(e)
        
        return doc_result
    
    def _replace_text_with_comment(self, doc: Document, old_text: str, new_text: str, comment: str, severity: str) -> bool:
        """Replace text and add comment with severity highlighting"""
        
        found = False
        severity_colors = {
            "CRITICAL": RGBColor(255, 0, 0),    
            "HIGH": RGBColor(255, 165, 0),     
            "MEDIUM": RGBColor(255, 255, 0),    
            "LOW": RGBColor(0, 255, 0)          
        }
        
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                
                comment_para = doc.add_paragraph()
                comment_run = comment_para.add_run(f"[{severity} COMPLIANCE FIX] {comment}")
                comment_run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))
                comment_run.bold = True
                
                found = True
                break
        
        return found
    
    def _add_compliance_header(self, doc: Document, total_edits: int, filename: str):
        """Add compliance review header to document"""
        
        
        header_para = doc.paragraphs[0]._element
        new_para = doc.add_paragraph()
        
       
        header_para.getparent().insert(0, new_para._element)
        
        new_para.text = f"ADGM COMPLIANCE REVIEW - {filename}\nTotal Edits Made: {total_edits}\nReview Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n" + "="*50 + "\n"
        new_para.runs[0].bold = True
        new_para.runs[0].font.color.rgb = RGBColor(0, 0, 255)
